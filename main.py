import argparse
import json
import os
from typing import Optional

from docx import Document
from dotenv import load_dotenv
from langchain.output_parsers import PydanticOutputParser
from langchain_openai import ChatOpenAI
from langgraph.prebuilt import create_react_agent

from models import ProfessionalProfile
from render_markdown import create_resume_markdown
from tools import ask_question


def read_docx_file(file_path: str) -> str:
    """Extract raw text from DOCX file."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    doc = Document(file_path)
    raw_text = ""

    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            raw_text += paragraph.text + "\n"

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    raw_text += cell.text + " "
            raw_text += "\n"

    if not raw_text.strip():
        raise ValueError("No text could be extracted from the DOCX file")

    return raw_text.strip()


class ResumeBuilderAgent:
    def __init__(self):
        # Set up the parser for structured output
        parser = PydanticOutputParser(pydantic_object=ProfessionalProfile)

        # System prompt for the React Agent
        self._prompt = (
            "You are a Senior Resume Writer acting on behalf of the candidate. "
            "Your job is to transform raw CV text into a clean, ATS-friendly structured professional profile.\n\n"

            "You will be given:\n"
            "- Raw text extracted from a candidate’s CV\n"
            "- An optional job description (JD) for context\n\n"

            "Your task:\n"
            "1. Analyze the CV text and create a structured professional profile in JSON format (per the parser schema).\n"
            "2. If a job description is provided, adapt the profile to highlight relevant skills and tailor the summary.\n"
            "3. If critical information is missing or ambiguous, use ask_question to request clarification (one at a time).\n"
            "4. Return the complete structured profile once all information is gathered.\n\n"

            "IMPORTANT:\n"
            "- Always incorporate JD keywords and reorder content to match JD priorities.\n"
            "- Focus on KPIs and measurable impact (latency, cost, quality, releases). "
            "If statements are vague (e.g., 'worked on', 'responsible for'), ask for metrics/examples.\n"
            "- Keep bullets ≤ 2 lines, remove duplication, respect NDA, and never fabricate facts.\n"
            "- Always produce a complete, accurate, ATS-optimized profile.\n\n"

            "Work step by step and ensure each tool call uses the most up-to-date information.\n"
            f"{parser.get_format_instructions()}"
        )

        # Create the React Agent with tools
        self._agent = create_react_agent(
            model=ChatOpenAI(model="gpt-5"),
            tools=[
                ask_question
            ],
            prompt=self._prompt,
            response_format=ProfessionalProfile
        )

    def run(self, cv_text: str, job_description: Optional[str] = None) -> ProfessionalProfile:
        """
        Run the resume processing agent.
        
        Args:
            cv_text: Raw text extracted from reference CV
            job_description: Optional job description for context and adaptation
            
        Returns:
            ProfessionalProfile: Complete structured profile
        """

        # Prepare agent input
        task_prompt = f"""Analyze this resume and build a complete structured professional profile, 
        adapting it to the job description if provided and filling in missing details where possible:

        Resume text:
        {cv_text}"""

        if job_description:
            task_prompt += f"""

        Job Description (use for context and adaptation):
        {job_description}"""

        task_prompt += "\n\nAsk clarifying questions if needed."

        # Run the agent
        print("🤖 Processing with AI agent...")
        result = self._agent.invoke(
            input={"messages": [("human", task_prompt)]},
            print_mode="values",
            config={"configurable": {"thread_id": "resume_processing"}}
        )

        return result


def main():
    """
    Main CLI entry point.
    """
    parser = argparse.ArgumentParser(description="AI Resume Builder - Convert DOCX resume to clean Markdown format")
    parser.add_argument("--in", dest="input_file", required=True, help="Input DOCX file path")
    parser.add_argument("--out", dest="output_file", default="resume.md", help="Output Markdown file path")
    parser.add_argument("--jd", dest="jd_file", help="Optional job description text file for adaptation")

    args = parser.parse_args()

    load_dotenv()

    # Validate OpenAI API key
    if not os.getenv("OPENAI_API_KEY"):
        print("Error: OPENAI_API_KEY environment variable is required")
        print("Please set your OpenAI API key: export OPENAI_API_KEY='your-key-here'")
        return

    # Validate input file
    if not os.path.exists(args.input_file):
        print(f"Error: Input file not found: {args.input_file}")
        return

    # Load job description if provided
    job_description = read_docx_file(args.jd_file)

    print("🚀 Starting AI Resume Builder...")
    print(f"📄 Input: {args.input_file}")
    print(f"📝 Output: {args.output_file}")
    if job_description:
        print(f"🎯 Job Description: {args.jd_file}")
    print("-" * 50)

    # Step 1: Read reference CV
    print("📄 Reading reference CV...")
    cv_text = read_docx_file(args.input_file)

    print(f"CV Text \n {cv_text}")

    # Step 2: Run the AI agent
    agent = ResumeBuilderAgent()
    print("🤖 Starting AI Resume Builder...")

    output = agent.run(cv_text, job_description)

    messages = output["messages"]

    for msg in messages:
        print(msg.content)

    result_dict = json.loads(messages[-1].content)

    print(result_dict)
    profile = ProfessionalProfile.model_validate(result_dict)

    # Step 3: Create final resume
    print("📝 Creating final resume...")
    final_path = create_resume_markdown(profile, args.output_file)

    print("\\n✅ Resume processing completed!")
    print(f"📋 Output saved to: {final_path}")


if __name__ == "__main__":
    main()
